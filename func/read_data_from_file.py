import pandas as pd
from docx import Document
import glob
import os


def read_csv(temp_file, uploaded_file, file_name):
    """
    From csv file read data. Convert to pandas df and creates new column.
    Returns new column.

    Args:
        temp_file: file created in tep folder
        uploaded_file: file name whitch are uploaded
        file_name: uploaded file name
    """
    with open(temp_file, "wb") as f:
        f.write(uploaded_file.getvalue())

    df = pd.read_csv(f"data/temp_files/{file_name}")
    df["id_overview"] = df["id"].astype("str") + " " + df["overview"]

    df_id_overview = df[["id_overview"]]

    return df_id_overview


def read_excel(temp_file, uploaded_file, file_name):
    """
    From excel file read data. Convert to pandas df and creates new column.
    Returns new column.

    Args:
        temp_file: file created in tep folder
        uploaded_file: file name whitch are uploaded
        file_name: uploaded file name
    """
    with open(temp_file, "wb") as f:
        f.write(uploaded_file.getvalue())

    df = pd.read_excel(f"data/temp_files/{file_name}")
    df["id_overview"] = df["id"].astype("str") + " " + df["overview"]
    df_id_overview = df[["id_overview"]]

    return df_id_overview


def read_word(temp_file, uploaded_file, file_name):
    """
    From docx file read data. Convert to pandas df and creates new column.
    Returns new column.

    Args:
        temp_file: file created in tep folder
        uploaded_file: file name whitch are uploaded
        file_name: uploaded file name
    """
    with open(temp_file, "wb") as f:
        f.write(uploaded_file.getvalue())

    document = Document(f"data/temp_files/{file_name}")
    table = document.tables[0]
    df = pd.DataFrame([[cell.text for cell in row.cells] for row in table.rows])
    df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
    df["id_overview"] = df["id"].astype("str") + " " + df["overview"]
    df_id_overview = df[["id_overview"]]

    return df_id_overview


def get_recomendation_id(answer) -> pd.DataFrame:
    """
    Get answer and return filtered results by id
    """
    temp_folder = r"data/temp_files/"
    lst_file = glob.glob(os.path.join(temp_folder, "*"))
    temp_file = []
    for i in lst_file:
        if i.endswith("txt"):
            continue
        else:
            temp_file.append(i.split("\\")[-1])

    id_lst = []
    if temp_file[0].endswith(".csv"):
        df = pd.read_csv(f"data/temp_files/{temp_file[0]}")
        for id in range(0, len(answer)):
            id_lst.append(int(answer[id].page_content.split(" ")[0].strip('"')))

        df = df.loc[df["id"].isin(id_lst)]

        return df

    elif temp_file[0].endswith(".xlsx"):
        df = pd.read_excel(f"data/temp_files/{temp_file[0]}")
        for id in range(0, len(answer)):
            id_lst.append(int(answer[id].page_content.split(" ")[0].strip('"')))

        df = df.loc[df["id"].isin(id_lst)]

        return df

    else:
        document = Document(f"data/temp_files/{temp_file[0]}")
        table = document.tables[0]
        df = pd.DataFrame([[cell.text for cell in row.cells] for row in table.rows])
        df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
        for id in range(0, len(answer)):
            id_lst.append(int(answer[id].page_content.split(" ")[0].strip('"')))
        df["id"] = df["id"].astype("int")

        df = df.loc[df["id"].isin(id_lst)]

        return df
